import docx
document = docx.Document()

document.add_heading('Result', 0)
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Image'
hdr_cells[1].text = 'SIFT'
hdr_cells[2].text = 'SURF'

sift = {}
sift_file = open('sift.txt', 'r', encoding='utf-8')
for line in sift_file:
    x = line.strip().split('->')
    sift[x[0]] = x[1]

surf = {}
sift_file = open('surf.txt', 'r', encoding='utf-8')
for line in sift_file:
    x = line.strip().split('->')
    surf[x[0]] = x[1]

for key in sift.keys():
    try:
        row_cells = table.add_row().cells
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run()
        print('test/data/' + key)
        row_cells[1].text = sift[key]
        row_cells[2].text = surf[key]
        run.add_picture('test/data/' + key, width=1400000, height=1400000)
    except Exception:
        print('Wrong action')

document.add_page_break()

document.save('result.docx')